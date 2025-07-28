import fitz  # PyMuPDF
from sqlalchemy.orm import Session
from typing import List, Optional
import os
import docx
from pathlib import Path

from models import Document, DocumentEmbedding
from schemas import DocumentCreate

class DocumentService:
    """
    Service for document operations: CRUD, text extraction, file management.
    Does NOT handle embeddings - that's handled by EmbeddingService.
    """
    
    def __init__(self, db: Session):
        self.db = db

    def create_document(self, document_data: dict) -> Document:
        """Create a new document record"""
        db_document = Document(**document_data)
        self.db.add(db_document)
        self.db.commit()
        self.db.refresh(db_document)
        return db_document

    def get_document(self, document_id: int) -> Optional[Document]:
        """Get a document by ID"""
        return self.db.query(Document).filter(Document.id == document_id).first()

    def get_documents(self, skip: int = 0, limit: int = 100) -> List[Document]:
        """Get all documents with pagination"""
        return self.db.query(Document).offset(skip).limit(limit).all()

    def update_extracted_text(self, document_id: int, extracted_text: str) -> bool:
        """Update the extracted text for a document"""
        document = self.get_document(document_id)
        if document:
            document.extracted_text = extracted_text
            self.db.commit()
            return True
        return False

    def delete_document(self, document_id: int) -> bool:
        """Delete a document (embeddings will be cascade deleted by database)"""
        document = self.get_document(document_id)
        if document:
            # Note: Associated embeddings are deleted by CASCADE in database schema
            # or should be explicitly deleted by EmbeddingService before calling this
            self.db.delete(document)
            self.db.commit()
            return True
        return False

    def extract_text(self, file_path: str, content_type: str) -> str:
        """Extract text from various file formats"""
        try:
            if content_type == "application/pdf":
                return self._extract_pdf_text(file_path)
            elif content_type == "text/plain":
                return self._extract_txt_text(file_path)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_docx_text(file_path)
            elif content_type == "application/msword":
                # For older .doc files, you might need python-docx2txt or other libraries
                return self._extract_doc_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {content_type}")
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()

    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text).strip()

    def _extract_doc_text(self, file_path: str) -> str:
        """Extract text from DOC file (placeholder - requires additional library)"""
        # You might want to use python-docx2txt or convert to DOCX first
        raise NotImplementedError("DOC file extraction not implemented. Please convert to DOCX.")

    def get_document_metadata(self, document_id: int) -> dict:
        """Get document metadata including embedding status"""
        document = self.get_document(document_id)
        if not document:
            return None
            
        embeddings_count = self.db.query(DocumentEmbedding).filter(
            DocumentEmbedding.document_id == document_id
        ).count()
        
        return {
            "id": document.id,
            "filename": document.filename,
            "original_filename": document.original_filename,
            "file_size": document.file_size,
            "content_type": document.content_type,
            "has_extracted_text": bool(document.extracted_text),
            "text_length": len(document.extracted_text) if document.extracted_text else 0,
            "embeddings_count": embeddings_count,
            "embeddings_generated": embeddings_count > 0,
            "created_at": document.created_at,
            "updated_at": document.updated_at
        }

    def get_documents_summary(self) -> dict:
        """Get summary statistics for all documents"""
        total_documents = self.db.query(Document).count()
        documents_with_text = self.db.query(Document).filter(
            Document.extracted_text.isnot(None)
        ).count()
        
        # Count documents with embeddings
        documents_with_embeddings = self.db.query(Document).join(DocumentEmbedding).distinct().count()
        
        return {
            "total_documents": total_documents,
            "documents_with_text": documents_with_text,
            "documents_with_embeddings": documents_with_embeddings,
            "documents_pending_processing": total_documents - documents_with_text,
            "documents_ready_for_embeddings": documents_with_text - documents_with_embeddings
        }